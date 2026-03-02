import shutil
from pathlib import Path

from docx import Document

from redakt.core.entities import PIIEntity
from redakt.parsers.base import BaseParser, ParseResult


class DocxParser(BaseParser):

    def extract_text(self, file_path: Path) -> ParseResult:
        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        text = "\n".join(parts)
        return ParseResult(
            text=text,
            file_path=file_path,
            metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
        )

    def create_anonymized(
        self,
        file_path: Path,
        entities: list[PIIEntity],
        output_path: Path | None = None,
    ) -> Path:
        output_path = output_path or self.output_filename(file_path)
        shutil.copy2(file_path, output_path)

        doc = Document(str(output_path))
        replacements = {e.original: e.placeholder for e in entities}
        # Sort by length descending to prevent partial matches
        sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

        # Replace in paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, sorted_replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, sorted_replacements)

        # Replace in headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer is not None:
                    for para in header_footer.paragraphs:
                        self._replace_in_paragraph(para, sorted_replacements)

        doc.save(str(output_path))
        return output_path

    @staticmethod
    def _replace_in_paragraph(para, sorted_replacements: list[tuple[str, str]]):
        """Replace PII in a paragraph, handling text split across runs."""
        full_text = para.text
        needs_replacement = any(orig in full_text for orig, _ in sorted_replacements)
        if not needs_replacement:
            return

        # Concatenate all runs, perform replacement, put result in first run
        runs = para.runs
        if not runs:
            return

        combined = "".join(r.text for r in runs)
        for orig, placeholder in sorted_replacements:
            combined = combined.replace(orig, placeholder)

        # Preserve first run's formatting, clear the rest
        runs[0].text = combined
        for r in runs[1:]:
            r.text = ""
