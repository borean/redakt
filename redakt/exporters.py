"""Export redacted documents in multiple formats.

All exporters take the original text and entity spans, and produce a new
file with PII replaced by solid block characters (visual black bars).
"""

import html as _html
import re as _re
from pathlib import Path

from redakt.core.entities import PIIEntity
from redakt.core.redactor import TextSpan, render_redacted_plain


# ── TXT ──────────────────────────────────────────────────────────────────


def export_txt(text: str, spans: list[TextSpan], output_path: Path, *, age_mode: bool = False) -> Path:
    """Plain-text with unicode block bars."""
    redacted = render_redacted_plain(text, spans, age_mode=age_mode)
    output_path.write_text(redacted, encoding="utf-8")
    return output_path


# ── Markdown ─────────────────────────────────────────────────────────────


def export_md(text: str, spans: list[TextSpan], output_path: Path, *, age_mode: bool = False) -> Path:
    """Markdown document with block bars and a PII summary table."""
    redacted = render_redacted_plain(text, spans, age_mode=age_mode)
    lines = [
        "# Redacted Document\n",
        "---\n",
        redacted,
        "\n---\n",
        f"\n*{len(spans)} PII items redacted.*\n",
    ]
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


# ── PDF (via PyMuPDF Story API) ──────────────────────────────────────────


def export_pdf(text: str, spans: list[TextSpan], output_path: Path, *, age_mode: bool = False) -> Path:
    """PDF with block-bar redactions using PyMuPDF's HTML Story layout."""
    import fitz

    redacted = render_redacted_plain(text, spans, age_mode=age_mode)
    escaped = _html.escape(redacted).replace("\n", "<br>")

    html_doc = (
        "<html><body>"
        '<p style="font-family: sans-serif; font-size: 10pt; line-height: 1.6;">'
        f"{escaped}"
        "</p></body></html>"
    )

    try:
        # Story API: proper pagination + Unicode font support
        story = fitz.Story(html_doc)
        writer = fitz.DocumentWriter(str(output_path))

        page_rect = fitz.paper_rect("a4")
        content_rect = page_rect + (50, 50, -50, -50)  # 50pt margins

        more = True
        while more:
            dev = writer.begin_page(page_rect)
            more, _ = story.place(content_rect)
            story.draw(dev)
            writer.end_page()

        writer.close()
    except Exception:
        # Fallback for older PyMuPDF without Story API
        doc = fitz.open()
        page = doc.new_page(width=595.28, height=841.89)
        rect = fitz.Rect(50, 50, 545, 790)
        page.insert_textbox(
            rect,
            redacted,
            fontsize=10,
            fontname="helv",
            align=0,
        )
        doc.save(str(output_path))
        doc.close()

    return output_path


# ── DOCX ─────────────────────────────────────────────────────────────────


def export_docx(text: str, spans: list[TextSpan], output_path: Path, *, age_mode: bool = False) -> Path:
    """DOCX with proper black-bar shading over redacted text."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Process text line by line, splitting into normal + redacted runs
    lines = text.split("\n")
    char_offset = 0
    span_idx = 0

    for line in lines:
        para = doc.add_paragraph()
        line_start = char_offset
        line_end = char_offset + len(line)
        pos = line_start

        # Process spans that overlap this line
        while span_idx < len(spans) and spans[span_idx].start < line_end:
            span = spans[span_idx]
            if span.end <= line_start:
                span_idx += 1
                continue
            if span.start >= line_end:
                break

            # Normal text before the span
            if span.start > pos:
                run = para.add_run(text[pos : span.start])
                run.font.size = Pt(10)
                run.font.name = "Calibri"

            ph = span.entity.placeholder
            if age_mode and span.entity.category == "date" and not ph.startswith("["):
                # Age-converted date → show the age text
                run = para.add_run(ph)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            else:
                # Redacted run: black text on black background = invisible
                bar_text = "\u2588" * (span.end - span.start)
                run = para.add_run(bar_text)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(0, 0, 0)
                # Add black background shading
                rPr = run._r.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "000000")
                rPr.append(shd)

            pos = span.end
            span_idx += 1

        # Remaining normal text on this line
        if pos < line_end:
            run = para.add_run(text[pos:line_end])
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        char_offset = line_end + 1  # +1 for the \n

    doc.save(str(output_path))
    return output_path


# ── Dispatcher ───────────────────────────────────────────────────────────

EXPORT_FORMATS = {
    "PDF": (".pdf", export_pdf),
    "DOCX": (".docx", export_docx),
    "TXT": (".txt", export_txt),
    "Markdown": (".md", export_md),
}


def export_redacted(
    fmt: str,
    text: str,
    spans: list[TextSpan],
    input_path: Path,
    *,
    age_mode: bool = False,
) -> Path:
    """Export a redacted document in the given format.

    Returns the path to the created file.
    """
    ext, exporter = EXPORT_FORMATS[fmt]
    stem = input_path.stem
    output_path = input_path.parent / f"{stem}_redacted{ext}"
    return exporter(text, spans, output_path, age_mode=age_mode)
