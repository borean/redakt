"""Export redacted documents in multiple formats.

All exporters take the original text and entity spans, and produce a new
file with PII replaced by solid block characters (visual black bars).
"""

import html as _html
import re as _re
from pathlib import Path

from qwenkk.core.entities import PIIEntity
from qwenkk.core.redactor import TextSpan, render_redacted_plain


# ── TXT ──────────────────────────────────────────────────────────────────


def export_txt(text: str, spans: list[TextSpan], output_path: Path) -> Path:
    """Plain-text with unicode block bars."""
    redacted = render_redacted_plain(text, spans)
    output_path.write_text(redacted, encoding="utf-8")
    return output_path


# ── Markdown ─────────────────────────────────────────────────────────────


def export_md(text: str, spans: list[TextSpan], output_path: Path) -> Path:
    """Markdown document with block bars and a PII summary table."""
    redacted = render_redacted_plain(text, spans)
    lines = [
        "# Redacted Document\n",
        "---\n",
        redacted,
        "\n---\n",
        f"\n*{len(spans)} PII items redacted.*\n",
    ]
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


# ── PDF (via PyMuPDF Story API) ──────────────────────────────────────────


def export_pdf(text: str, spans: list[TextSpan], output_path: Path) -> Path:
    """PDF with block-bar redactions using PyMuPDF's HTML Story layout."""
    import fitz

    redacted = render_redacted_plain(text, spans)
    escaped = _html.escape(redacted).replace("\n", "<br>")

    html_doc = (
        "<html><body>"
        '<p style="font-family: sans-serif; font-size: 10pt; line-height: 1.6;">'
        f"{escaped}"
        "</p></body></html>"
    )

    try:
        # Story API: proper pagination + Unicode font support
        story = fitz.Story(html_doc)
        writer = fitz.DocumentWriter(str(output_path))

        page_rect = fitz.paper_rect("a4")
        content_rect = page_rect + (50, 50, -50, -50)  # 50pt margins

        more = True
        while more:
            dev = writer.begin_page(page_rect)
            more, _ = story.place(content_rect)
            story.draw(dev)
            writer.end_page()

        writer.close()
    except Exception:
        # Fallback for older PyMuPDF without Story API
        doc = fitz.open()
        page = doc.new_page(width=595.28, height=841.89)
        rect = fitz.Rect(50, 50, 545, 790)
        page.insert_textbox(
            rect,
            redacted,
            fontsize=10,
            fontname="helv",
            align=0,
        )
        doc.save(str(output_path))
        doc.close()

    return output_path


# ── DOCX ─────────────────────────────────────────────────────────────────


def export_docx(text: str, spans: list[TextSpan], output_path: Path) -> Path:
    """DOCX with proper black-bar shading over redacted text."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Process text line by line, splitting into normal + redacted runs
    lines = text.split("\n")
    char_offset = 0
    span_idx = 0

    for line in lines:
        para = doc.add_paragraph()
        line_start = char_offset
        line_end = char_offset + len(line)
        pos = line_start

        # Process spans that overlap this line
        while span_idx < len(spans) and spans[span_idx].start < line_end:
            span = spans[span_idx]
            if span.end <= line_start:
                span_idx += 1
                continue
            if span.start >= line_end:
                break

            # Normal text before the span
            if span.start > pos:
                run = para.add_run(text[pos : span.start])
                run.font.size = Pt(10)
                run.font.name = "Calibri"

            # Redacted run: black text on black background = invisible
            bar_text = "\u2588" * (span.end - span.start)
            run = para.add_run(bar_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0, 0, 0)
            # Add black background shading
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "000000")
            rPr.append(shd)

            pos = span.end
            span_idx += 1

        # Remaining normal text on this line
        if pos < line_end:
            run = para.add_run(text[pos:line_end])
            run.font.size = Pt(10)
            run.font.name = "Calibri"

        char_offset = line_end + 1  # +1 for the \n

    doc.save(str(output_path))
    return output_path


# ── Summary / Report exporters ──────────────────────────────────────────


def export_summary_txt(summary: str, output_path: Path) -> Path:
    """Write a plain-text summary."""
    output_path.write_text(summary, encoding="utf-8")
    return output_path


def export_summary_md(summary: str, output_path: Path) -> Path:
    """Write a Markdown summary with a document header."""
    md = f"# Document Summary\n\n{summary}\n"
    output_path.write_text(md, encoding="utf-8")
    return output_path


def export_summary_pdf(summary: str, output_path: Path) -> Path:
    """PDF summary with styled markdown using PyMuPDF's HTML Story layout."""
    import fitz
    import markdown

    md = markdown.Markdown(
        extensions=["fenced_code", "tables", "sane_lists", "smarty"],
        output_format="html",
    )
    body_html = md.convert(summary)

    css = """
    <style>
    body {
        font-family: Helvetica, Arial, sans-serif;
        font-size: 10pt;
        line-height: 1.6;
        color: #1a1a1a;
    }
    h1 { font-size: 16pt; font-weight: bold; margin-top: 14pt; margin-bottom: 6pt;
         border-bottom: 1px solid #cccccc; padding-bottom: 4pt; }
    h2 { font-size: 14pt; font-weight: bold; margin-top: 12pt; margin-bottom: 4pt; }
    h3 { font-size: 12pt; font-weight: bold; margin-top: 10pt; margin-bottom: 4pt; }
    h4 { font-size: 11pt; font-weight: bold; margin-top: 8pt; margin-bottom: 4pt; }
    p { margin: 4pt 0; }
    strong, b { font-weight: bold; }
    em, i { font-style: italic; }
    ul, ol { margin: 4pt 0; padding-left: 18pt; }
    li { margin: 2pt 0; }
    code { font-family: Courier, monospace; font-size: 9pt;
           background-color: #f0f0f0; padding: 1pt 3pt; }
    pre { font-family: Courier, monospace; font-size: 9pt;
          background-color: #f0f0f0; padding: 8pt; margin: 6pt 0;
          border: 1px solid #cccccc; }
    blockquote { border-left: 3pt solid #cccccc; padding-left: 10pt;
                 margin: 6pt 0; color: #555555; }
    table { width: 100%; border-collapse: collapse; margin: 6pt 0; }
    th { font-weight: bold; text-align: left; padding: 4pt 8pt;
         border-bottom: 1px solid #999999; }
    td { padding: 4pt 8pt; border-bottom: 1px solid #dddddd; }
    </style>
    """

    html_doc = f"<html><head>{css}</head><body>{body_html}</body></html>"

    try:
        story = fitz.Story(html_doc)
        writer = fitz.DocumentWriter(str(output_path))
        page_rect = fitz.paper_rect("a4")
        content_rect = page_rect + (50, 50, -50, -50)
        more = True
        while more:
            dev = writer.begin_page(page_rect)
            more, _ = story.place(content_rect)
            story.draw(dev)
            writer.end_page()
        writer.close()
    except Exception:
        doc = fitz.open()
        page = doc.new_page(width=595.28, height=841.89)
        rect = fitz.Rect(50, 50, 545, 790)
        page.insert_textbox(rect, summary, fontsize=10, fontname="helv", align=0)
        doc.save(str(output_path))
        doc.close()

    return output_path


def _strip_inline_md(text: str) -> str:
    """Remove inline markdown formatting for plain text contexts."""
    text = _re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = _re.sub(r"\*(.+?)\*", r"\1", text)
    text = _re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_inline_runs(para, text: str, default_size=10):
    """Parse inline markdown (bold, italic, code) and add as formatted runs."""
    from docx.shared import Pt

    pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
    parts = _re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            run = para.add_run(part)
        run.font.size = Pt(default_size)
        if not run.font.name:
            run.font.name = "Calibri"


def _render_markdown_to_docx(doc, summary: str):
    """Parse markdown text and add styled elements to a python-docx Document."""
    from docx.shared import Pt, RGBColor

    lines = summary.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        heading_match = _re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = _strip_inline_md(heading_match.group(2).strip())
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Unordered list: - item or * item (but not **bold**)
        if _re.match(r"^[-*]\s+", stripped) and not stripped.startswith("**"):
            text = _re.sub(r"^[-*]\s+", "", stripped)
            try:
                para = doc.add_paragraph(style="List Bullet")
            except KeyError:
                para = doc.add_paragraph()
                para.add_run("\u2022 ").bold = False
            _add_inline_runs(para, text)
            i += 1
            continue

        # Ordered list: 1. item
        if _re.match(r"^\d+\.\s+", stripped):
            text = _re.sub(r"^\d+\.\s+", "", stripped)
            try:
                para = doc.add_paragraph(style="List Number")
            except KeyError:
                para = doc.add_paragraph()
            _add_inline_runs(para, text)
            i += 1
            continue

        # Horizontal rule
        if _re.match(r"^[-*_]{3,}$", stripped):
            para = doc.add_paragraph()
            run = para.add_run("\u2500" * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(8)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        _add_inline_runs(para, stripped)
        i += 1


def export_summary_docx(summary: str, output_path: Path) -> Path:
    """DOCX summary with proper markdown formatting."""
    from docx import Document

    doc = Document()
    doc.add_heading("Document Summary", level=1)
    _render_markdown_to_docx(doc, summary)
    doc.save(str(output_path))
    return output_path


# ── Dispatcher ───────────────────────────────────────────────────────────

EXPORT_FORMATS = {
    "PDF": (".pdf", export_pdf),
    "DOCX": (".docx", export_docx),
    "TXT": (".txt", export_txt),
    "Markdown": (".md", export_md),
}


def export_redacted(
    fmt: str,
    text: str,
    spans: list[TextSpan],
    input_path: Path,
) -> Path:
    """Export a redacted document in the given format.

    Returns the path to the created file.
    """
    ext, exporter = EXPORT_FORMATS[fmt]
    stem = input_path.stem
    output_path = input_path.parent / f"{stem}_redacted{ext}"
    return exporter(text, spans, output_path)


SUMMARY_EXPORT_FORMATS = {
    "PDF": (".pdf", export_summary_pdf),
    "DOCX": (".docx", export_summary_docx),
    "TXT": (".txt", export_summary_txt),
    "Markdown": (".md", export_summary_md),
}


def export_summary(fmt: str, summary: str, input_path: Path) -> Path:
    """Export a document summary in the given format."""
    ext, exporter = SUMMARY_EXPORT_FORMATS[fmt]
    stem = input_path.stem
    output_path = input_path.parent / f"{stem}_summary{ext}"
    return exporter(summary, output_path)
